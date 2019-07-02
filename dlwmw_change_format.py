import docx
import os
from docx import Document
from docx.shared import Inches
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.text import WD_LINE_SPACING
path_docx='C://Users//Peter//Desktop//dlwmw-docx//'
docxlist=os.listdir(path_docx)
for n in docxlist:
    name=os.path.join(path_docx+n)
    doc=Document(name)
    print(doc.paragraphs[0].text)
    doc.paragraphs[0].paragraph_format.alignment=WD_ALIGN_PARAGRAPH.CENTER
    doc.paragraphs[0].style='Normal'
    doc.paragraphs[0].style.font.name='宋体'
    doc.paragraphs[0].style.font.size=Pt(18)
    doc.paragraphs[0].paragraph_format.space_before=Pt(9)
    doc.paragraphs[0].paragraph_format.space_after=Pt(9)
    doc.paragraphs[1].style.font.name='宋体'
    doc.paragraphs[1].style.font.size=Pt(14)
    for p in doc.paragraphs[1:]:
        p.paragraph_format.line_spacing_rule=WD_LINE_SPACING.MULTIPLE
        p.paragraph_format.line_spacing=1.5
        p.paragraph_format.space_before=0
        p.paragraph_format.space_after=0
        p.paragraph_format.first_line_indent=Pt(28)
    doc.save(name)
